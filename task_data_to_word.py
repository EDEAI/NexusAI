from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from typing import Dict, Any
import os
import uuid

def main(data: Dict[str, Any], file_name: str = None) -> dict:
    """
    根据传入的数据生成Word文档
    格式：{name, content, table/tables, image/images, children}
    
    Args:
        data: 数据字典，包含文档内容
        file_name: 文件名（不含扩展名），如果为None则使用UUID生成
    
    Returns:
        dict: 包含操作结果和文件路径的字典
    """
    # # Generate a unique filename using UUID
    # if file_name is None:
    #     file_name = f"{uuid.uuid4()}.docx"
    # else:
    #     # 确保文件名有.docx扩展名
    #     if not file_name.endswith('.docx'):
    #         file_name = f"{file_name}.docx"
    # file_path = f"/storage/{file_name}"
    
    # # Ensure directory exists
    # directory = os.path.dirname(file_path)
    # if not os.path.exists(directory):
    #     os.makedirs(directory, exist_ok=True)


    # Generate filename
    if file_name is None:
        file_name = f"{uuid.uuid4()}.docx"
    else:
        # 确保文件名有.docx扩展名
        if not file_name.endswith('.docx'):
            file_name = f"{file_name}.docx"
    # 使用相对路径的storage文件夹
    output_dir = "storage"
    file_path = os.path.join(output_dir, file_name)
    
    # Ensure directory exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    # Create new Word document
    doc = Document()
    
    # Set Chinese font
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # Add document title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = '黑体'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_before = Pt(24)
    title_style.paragraph_format.space_after = Pt(18)
    # Set Chinese font
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # Add level counters
    level_counters = {}
    
    # Chinese numbers configuration (unified)
    CHINESE_NUMBERS = [
        '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
        '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十',
        '二十一', '二十二', '二十三', '二十四', '二十五', '二十六', '二十七', '二十八', '二十九', '三十',
        '三十一', '三十二', '三十三', '三十四', '三十五', '三十六', '三十七', '三十八', '三十九', '四十',
        '四十一', '四十二', '四十三', '四十四', '四十五', '四十六', '四十七', '四十八', '四十九', '五十'
    ]
    
    # Generate Chinese number prefixes for pattern matching
    CHINESE_NUMBER_PREFIXES = tuple(num + '、' for num in CHINESE_NUMBERS)
    
    def reset_counters_from_level(level):
        """Reset counters from specified level and below"""
        keys_to_remove = [k for k in level_counters.keys() if k >= level]
        for key in keys_to_remove:
            del level_counters[key]
    
    def get_level_number(level):
        """Get level number for specified level"""
        # Reset deeper level counters
        reset_counters_from_level(level + 1)
        
        # Increment current level counter
        if level not in level_counters:
            level_counters[level] = 0
        level_counters[level] += 1
        
        if level == 1:
            # First level uses Chinese numbers
            if level_counters[level] <= len(CHINESE_NUMBERS):
                return CHINESE_NUMBERS[level_counters[level] - 1] + '、'
            else:
                return f"{level_counters[level]}、"
        elif level == 2:
            # Second level uses "1.1, 1.2" format, corresponding to first level number
            first_level_num = level_counters.get(1, 1)  # Get first level number
            return f"{first_level_num}.{level_counters[level]}、"
        else:
            # Third level and beyond use multi-level numbers
            if level == 3:
                first_level_num = level_counters.get(1, 1)
                return f"{first_level_num}.{level_counters[2]}.{level_counters[level]} "
            elif level == 4:
                first_level_num = level_counters.get(1, 1)
                return f"{first_level_num}.{level_counters[2]}.{level_counters[3]}.{level_counters[level]} "
            elif level == 5:
                first_level_num = level_counters.get(1, 1)
                return f"{first_level_num}.{level_counters[2]}.{level_counters[3]}.{level_counters[4]}.{level_counters[level]} "
            else:
                # Deeper levels continue expansion
                first_level_num = level_counters.get(1, 1)
                parts = [str(first_level_num)]
                for i in range(2, level + 1):
                    if i in level_counters:
                        parts.append(str(level_counters[i]))
                return '.'.join(parts) + ' '

    # Internal function: add heading with level numbering
    def add_heading(text, level=1):
        # Get level number
        level_number = get_level_number(level)
        numbered_text = level_number + text
        
        # Define style configuration for each level
        level_configs = {
            1: {'font_size': 16, 'space_before': 18, 'space_after': 12},
            2: {'font_size': 14, 'space_before': 12, 'space_after': 8},
            3: {'font_size': 13, 'space_before': 10, 'space_after': 6},
            # Default configuration for level 4 and below
            'default': {'font_size': 12, 'space_before': 8, 'space_after': 4}
        }
        
        def apply_heading_style(heading, font_size, space_before, space_after):
            """Common function to apply heading styles"""
            heading.style.font.name = '黑体'
            heading.style.font.size = Pt(font_size)
            heading.style.font.bold = True
            heading.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading.style.paragraph_format.space_before = Pt(space_before)
            heading.style.paragraph_format.space_after = Pt(space_after)
            # Set Chinese font to ensure proper display
            heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        if level <= 3:
            # Level 1-3 headings use Word built-in heading styles
            heading = doc.add_heading(numbered_text, level=level)
            config = level_configs[level]
            apply_heading_style(heading, config['font_size'], config['space_before'], config['space_after'])
            return heading
        else:
            # Level 4+ headings also use add_heading but with level=3 style (Word supports up to level 3)
            heading = doc.add_heading(numbered_text, level=3)
            config = level_configs['default']
            apply_heading_style(heading, config['font_size'], config['space_before'], config['space_after'])
            return heading
    
    # 处理内容的函数
    def add_content(content_text):
        """添加内容段落"""
        if content_text and content_text.strip():
            p = doc.add_paragraph(content_text.strip())
            p.paragraph_format.space_after = Pt(6)
    
    # 处理表格的函数
    def add_table(table_data):
        """添加单个表格"""
        if not table_data:
            return
        
        # 支持新格式：{headers: [...], rows: [[...], [...]]}
        if isinstance(table_data, dict) and 'headers' in table_data and 'rows' in table_data:
            headers = table_data['headers']
            rows = table_data['rows']
            
            if not headers or not rows:
                return
            
            # 创建表格
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # 添加表头
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = str(header)
                # 设置表头样式
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '黑体'
            
            # 添加数据行
            for row_data in rows:
                row_cells = table.add_row().cells
                for i, cell_value in enumerate(row_data):
                    if i < len(row_cells):  # 防止索引越界
                        row_cells[i].text = str(cell_value) if cell_value is not None else ''
        
        # 兼容旧格式：[{key: value}, {key: value}]
        elif isinstance(table_data, list) and len(table_data) > 0:
            # 获取表头（第一行的键）
            headers = list(table_data[0].keys()) if isinstance(table_data[0], dict) else []
            
            if not headers:
                return
            
            # 创建表格
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # 添加表头
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = str(header)
                # 设置表头样式
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = '黑体'
            
            # 添加数据行
            for row_data in table_data:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    cell_value = row_data.get(header, '')
                    row_cells[i].text = str(cell_value) if cell_value is not None else ''
        
        else:
            return
        
        # 添加表格后的间距
        doc.add_paragraph()
    
    # 处理多个表格的函数
    def add_tables(tables_data):
        """添加多个表格"""
        if not tables_data:
            return
        
        if isinstance(tables_data, list):
            # 多个表格
            for table_data in tables_data:
                add_table(table_data)
        else:
            # 单个表格（向后兼容）
            add_table(tables_data)
    
    # 处理图片的函数
    def add_image(image_info):
        """添加单个图片"""
        if not image_info or not isinstance(image_info, dict):
            return
        
        image_path = image_info.get('path', '')
        image_caption = image_info.get('caption', '')
        
        if not image_path:
            return
        
        try:
            # 检查是否是HTTP/HTTPS链接
            if image_path.startswith(('http://', 'https://')):
                # 对于HTTP图片，下载到临时文件
                import requests
                import tempfile
                
                try:
                    print(f"正在下载图片: {image_path}")
                    
                    # 添加请求头，模拟浏览器
                    headers = {
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                    }
                    
                    response = requests.get(image_path, timeout=10, headers=headers)
                    response.raise_for_status()
                    
                    # 检查内容
                    content_length = len(response.content)
                    content_type = response.headers.get('content-type', '').lower()
                    print(f"图片下载成功，大小: {content_length} 字节，类型: {content_type}")
                    
                    if content_length == 0:
                        raise Exception("下载的图片内容为空")
                    
                    # 根据内容类型确定文件扩展名
                    if 'png' in content_type:
                        suffix = '.png'
                    elif 'jpeg' in content_type or 'jpg' in content_type:
                        suffix = '.jpg'
                    elif 'gif' in content_type:
                        suffix = '.gif'
                    elif 'webp' in content_type:
                        suffix = '.webp'
                    else:
                        suffix = '.png'  # 默认使用png
                    
                    # 创建临时文件
                    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                        temp_file.write(response.content)
                        temp_path = temp_file.name
                    
                    print(f"临时文件创建: {temp_path}")
                    
                    # 验证文件是否创建成功
                    if not os.path.exists(temp_path):
                        raise Exception("临时文件创建失败")
                    
                    file_size = os.path.getsize(temp_path)
                    if file_size == 0:
                        raise Exception("临时文件为空")
                    
                    print(f"临时文件大小: {file_size} 字节")
                    
                    # 添加图片
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(temp_path, width=Inches(4))  # 设置图片宽度为4英寸
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    print("图片添加到文档成功")
                    
                    # 清理临时文件
                    os.unlink(temp_path)
                    
                except requests.exceptions.Timeout:
                    print("图片下载超时")
                    placeholder_p = doc.add_paragraph(f"[网络图片下载超时: {image_path}]")
                    placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in placeholder_p.runs:
                        run.font.italic = True
                        
                except requests.exceptions.ConnectionError:
                    print("图片下载连接错误")
                    placeholder_p = doc.add_paragraph(f"[网络图片连接失败: {image_path}]")
                    placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in placeholder_p.runs:
                        run.font.italic = True
                        
                except requests.exceptions.HTTPError as e:
                    print(f"图片下载HTTP错误: {e}")
                    placeholder_p = doc.add_paragraph(f"[网络图片HTTP错误: {image_path}]\n状态码: {e.response.status_code}")
                    placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in placeholder_p.runs:
                        run.font.italic = True
                        
                except Exception as e:
                    # HTTP图片下载失败，尝试备用图片
                    print(f"图片处理失败: {str(e)}")
                    
                    # 尝试使用备用的公共图片
                    backup_url = "https://via.placeholder.com/400x300.png?text=图片加载失败"
                    try:
                        print(f"尝试备用图片: {backup_url}")
                        backup_response = requests.get(backup_url, timeout=5)
                        backup_response.raise_for_status()
                        
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                            temp_file.write(backup_response.content)
                            temp_path = temp_file.name
                        
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(temp_path, width=Inches(4))
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        os.unlink(temp_path)
                        print("备用图片加载成功")
                        
                    except Exception as backup_e:
                        print(f"备用图片也失败: {backup_e}")
                        placeholder_p = doc.add_paragraph(f"[网络图片处理失败: {image_path}]\n原始错误: {str(e)}\n备用图片错误: {str(backup_e)}")
                        placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in placeholder_p.runs:
                            run.font.italic = True
                        
            elif os.path.exists(image_path):
                # 本地图片文件
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(4))  # 设置图片宽度为4英寸
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            else:
                # 如果图片不存在，添加占位符
                placeholder_p = doc.add_paragraph(f"[图片占位符: {image_path}]")
                placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in placeholder_p.runs:
                    run.font.italic = True
            
            # 添加图片说明
            if image_caption:
                caption_p = doc.add_paragraph(image_caption)
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption_p.runs:
                    run.font.size = Pt(10)
                    run.font.italic = True
            
            # 添加间距
            doc.add_paragraph()
            
        except Exception as e:
            # 添加错误信息
            error_p = doc.add_paragraph(f"[图片加载失败: {image_path}]")
            error_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in error_p.runs:
                run.font.italic = True
    
    # 处理多个图片的函数
    def add_images(images_data):
        """添加多个图片"""
        if not images_data:
            return
        
        if isinstance(images_data, list):
            # 多个图片
            for image_info in images_data:
                add_image(image_info)
        else:
            # 单个图片（向后兼容）
            add_image(images_data)
    
    # 递归处理children的函数
    def process_children(children, level=1):
        """递归处理子分类"""
        for category in children:
            # 添加标题
            if 'name' in category:
                add_heading(category['name'], level=level)
            
            # 添加内容
            if 'content' in category:
                add_content(category['content'])
            
            # 添加表格（支持单个和多个）
            if 'table' in category:
                add_table(category['table'])
            if 'tables' in category:
                add_tables(category['tables'])
            
            # 添加图片（支持单个和多个）
            if 'image' in category:
                add_image(category['image'])
            if 'images' in category:
                add_images(category['images'])
            
            # 递归处理子分类
            if 'children' in category and category['children']:
                process_children(category['children'], level=level+1)
    
    # 添加文档标题
    if 'name' in data:
        title = doc.add_paragraph(data['name'], style='CustomTitle')
    else:
        title = doc.add_paragraph('文档标题', style='CustomTitle')
    
    # 添加主要内容
    if 'content' in data:
        add_content(data['content'])
    
    # 添加主文档的表格（支持单个和多个）
    if 'table' in data:
        add_table(data['table'])
    if 'tables' in data:
        add_tables(data['tables'])
    
    # 添加主文档的图片（支持单个和多个）
    if 'image' in data:
        add_image(data['image'])
    if 'images' in data:
        add_images(data['images'])
    
    # 处理子分类
    if 'children' in data:
        process_children(data['children'], level=1)
    
    # Save document
    try:
        doc.save(file_path)
        return {
            "success": "True",
            "file_path": f"file://{file_path}"
        }
    except Exception as e:
        return {
            "success": "False",
            "file_path": ""
        }

# 测试函数
if __name__ == "__main__":
    # 导入测试数据
    from task_data import data
    
    # 直接运行测试
    result = main(data, "test_word_document")
    print("生成结果:", result)